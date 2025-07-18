# ==============================================================================
# RSOC_OS — Operational Support Suite for the Flex Regional Security Operations Center
#
# Copyright (c) 2025 Morgan Small
# All rights reserved.
#
# Permission is granted to current Flex RSOC personnel to use this software
# solely for official operational support and task automation.
#
# Use of this suite is implicitly permitted only while Morgan Small is employed
# within the Flex RSOC organizational structure. Should he be demoted,
# terminated, or otherwise removed from the RSOC hierarchy in any way,
# this implicit permission is revoked. Continued use of RSOC_OS following such
# circumstances is prohibited unless explicitly authorized by the original author.
#
# This program is intended to assist RSOC operators and supervisors in completing
# repetitive daily tasks efficiently and consistently. It is not designed to
# replace human oversight or operator judgment. RSOC personnel are still required
# to provide appropriate input, review outputs, and confirm that all
# generated content is accurate and appropriate for operational use.
#
# Redistribution:
# Redistribution, reproduction, or reuse of this software or any of its components
# outside the Flex RSOC environment is strictly prohibited without explicit,
# written permission from the author, Morgan Small.
#
# Attribution:
# Any derivative works, extensions, or adaptations of this software must
# include clear attribution to the original author, Morgan Small.
#
# External Dependencies:
# This software relies on third-party packages. Compatibility with future versions
# of those libraries is not guaranteed. It is the user's responsibility to maintain
# a stable environment for proper functionality.
#
# Disclaimer of Warranty:
# This software is provided "as is" without warranty of any kind, express or implied.
# In no event shall the author be held liable for any damages or losses arising
# from the use, misuse, or inability to use this software.
#
# Confidentiality:
# Portions of this software may contain proprietary logic or access confidential
# systems and workflows. Users are expected to treat the internal logic, file paths,
# and associated data structures as confidential and not disclose them outside of
# authorized RSOC personnel.
#
# Version Integrity:
# Modifications to this software should be version-controlled and approved by the
# original author. Unauthorized edits or forks may compromise the tool's intended
# functionality and are strongly discouraged.
#
# Contact:
# For support, feedback, or licensing inquiries, contact:
# Morgan Small — morgan.small@flex.com OR jamiesmall0718@gmail.com
# ==============================================================================

import os
import json
from datetime import datetime, timedelta
from PySide6.QtWidgets import (
    QWidget, QLabel, QLineEdit, QComboBox, QPushButton,
    QVBoxLayout, QHBoxLayout, QMessageBox
)
from PySide6.QtCore import Qt
from docx import Document # type: ignore

BOLO_TEMPLATE_PATH = r"C:\Users\User\Box\RSOC\BOLO\BOLO Subject Name Date - Template.docx"
BOLO_BASE_DIR = os.path.expandvars(r"C:\Users\%USERNAME%\Box\RSOC\BOLO")
BOLO_COUNTER_FILE = os.path.expandvars(r"C:\Users\%USERNAME%\Box\RSOC\Officer Files\M. Small\RSOC_OS Shared Files\bolo_counter_log.txt")

REGIONAL_SITES = {
    "BFG": "Buffalo Grove, IL",
    "CLB": "West Columbia, SC",
    "CVL": "Coopersville, MI",
    "MCT": "Manchester, CT",
    "HOL": "Hollis, New Hampshire",
    "MEM": "Memphis, TN",
    "NMK": "Newmarket, Ontario, CA",
    "RHL": "Farmington Hills, MI",
    "RIC": "Richmond, VA",
    "SJC": "San Jose, CA",
    "STH": "Sterling Heights, MI",
    "SLC": "Salt Lake City, UT"
}

def get_next_bolo_serial():
    today_str = datetime.now().strftime("%m%d%y")
    counter = 1

    if not os.path.exists(BOLO_COUNTER_FILE):
        with open(BOLO_COUNTER_FILE, "w") as f:
            f.write(f"{today_str},1\n")
        return f"{today_str}.0001"

    with open(BOLO_COUNTER_FILE, "r+") as f:
        lines = f.readlines()
        if lines and lines[-1].startswith(today_str):
            counter = int(lines[-1].split(",")[1]) + 1
        f.write(f"{today_str},{counter}\n")

    return f"{today_str}.{str(counter).zfill(4)}"

def modify_bolo_document(file_path, serial, issue_date, end_date, subject_name):
    try:
        doc = Document(file_path)

        def replace_in_run(run, old, new):
            if old in run.text:
                run.text = run.text.replace(old, new)

        # Replace in paragraphs
        replaced_issue = False
        for para in doc.paragraphs:
            for run in para.runs:
                replace_in_run(run, "000000.0000", serial)
                if "00/00/0000" in run.text:
                    if not replaced_issue:
                        run.text = run.text.replace("00/00/0000", issue_date, 1)
                        replaced_issue = True
                    else:
                        run.text = run.text.replace("00/00/0000", end_date, 1)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            replace_in_run(run, "000000.0000", serial)
                            if "00/00/0000" in run.text:
                                if not replaced_issue:
                                    run.text = run.text.replace("00/00/0000", issue_date, 1)
                                    replaced_issue = True
                                else:
                                    run.text = run.text.replace("00/00/0000", end_date, 1)
        
        # Insert subject's name into the cell right of the one labeled "Name"
        for table in doc.tables:
            for row in table.rows:
                for idx, cell in enumerate(row.cells[:-1]):  # avoid index error
                    if cell.text.strip() == "Subject Name":
                        target_cell = row.cells[idx + 1]
                        if target_cell.paragraphs and target_cell.paragraphs[0].runs:
                            run_format = target_cell.paragraphs[0].runs[0]
                            target_cell.paragraphs[0].clear()  # Clear all text
                            new_run = target_cell.paragraphs[0].add_run(subject_name)
                            new_run.bold = run_format.bold
                            new_run.italic = run_format.italic
                            new_run.font.name = run_format.font.name
                            new_run.font.size = run_format.font.size
                        else:
                            target_cell.text = subject_name
                        # Break out of all loops
                        break  # exits row.cells loop
                else:
                    continue  # only reached if inner loop didn't break
                break  # exits table.rows loop
            else:
                continue
            break  # exits doc.tables loop

        doc.save(file_path)
        return True

    except Exception as e:
        print(f"[!] Word modification failed: {e}")
        return False

class BOLOGenerator(QWidget):
    def __init__(self):
        super().__init__()
        self.setObjectName("MainWidget")
        self.setLayout(QVBoxLayout())
        self.layout().setAlignment(Qt.AlignmentFlag.AlignTop)

        self.title = QLabel("BOLO Generator")
        self.title.setObjectName("SitrepTitle")

        self.first_name_input = QLineEdit()
        self.first_name_input.setPlaceholderText("First Name")

        self.last_name_input = QLineEdit()
        self.last_name_input.setPlaceholderText("Last Name")

        self.scope_selector = QComboBox()
        self.scope_selector.addItems(["Local", "Regional"])
        self.scope_selector.currentTextChanged.connect(self.toggle_site_dropdown)

        self.site_selector = QComboBox()
        self.site_selector.addItems(REGIONAL_SITES.keys())
        self.site_selector.setVisible(False)

        self.generate_button = QPushButton("Generate BOLO")
        self.generate_button.clicked.connect(self.generate_bolo)

        input_layout = QHBoxLayout()
        input_layout.addWidget(self.first_name_input)
        input_layout.addWidget(self.last_name_input)
        input_layout.addWidget(self.scope_selector)
        input_layout.addWidget(self.site_selector)
        input_layout.addWidget(self.generate_button)

        self.layout().addWidget(self.title)
        self.layout().addLayout(input_layout)

    def toggle_site_dropdown(self, value):
        self.site_selector.setVisible(value == "Regional")

    def generate_bolo(self):
        first = self.first_name_input.text().strip()
        last = self.last_name_input.text().strip()
        scope = self.scope_selector.currentText()
        site = self.site_selector.currentText() if scope == "Regional" else "ATX"

        if not first or not last:
            QMessageBox.warning(self, "Missing Info", "Please enter both first and last name.")
            return

        serial = get_next_bolo_serial()
        today = datetime.now()
        issue_date = today.strftime("%m/%d/%Y")
        end_date = (today + timedelta(days=30)).strftime("%m/%d/%Y")

        # Determine base folder based on scope
        base_folder = os.path.join(BOLO_BASE_DIR, "Regional" if scope == "Regional" else "Local - ATX")
        year_folder = f"BOLO {today.year}"
        full_folder_path = os.path.join(base_folder, year_folder)
        os.makedirs(full_folder_path, exist_ok=True)

        filename = f"{today.strftime('%m.%d.%Y')} BOLO - {first} {last}.docx"
        output_path = os.path.join(full_folder_path, filename)

        try:
            # Copy the template to the output path
            from shutil import copyfile
            copyfile(BOLO_TEMPLATE_PATH, output_path)

            # Modify the copied document
            success = modify_bolo_document(output_path, serial, issue_date, end_date, f"{first} {last}")
            if not success:
                raise Exception("Failed to modify Word document.")

            QMessageBox.information(self, "BOLO Created", f"BOLO generated at:\n{output_path}")
            os.startfile(output_path)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to generate BOLO:\n{str(e)}")


def get_bolo_generator_widget():
    return BOLOGenerator()
